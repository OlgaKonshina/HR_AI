import json
import re
from pathlib import Path
import pandas as pd
import docx
from striprtf.striprtf import rtf_to_text
from transformers import AutoTokenizer, AutoModel
import torch
import torch.nn.functional as F
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ PyMuPDF не доступен: {e}")
    FITZ_AVAILABLE = False

class DocumentReader:
    def __init__(self, filepath: str):
        self.filepath = Path(filepath)
        self.text = ""

    def extract_text(self):
        suffix = self.filepath.suffix.lower()

        if suffix == ".pdf":
            self.text = self._read_pdf()
        elif suffix == ".txt":
            self.text = self.filepath.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".json":
            self.text = self._read_json()
        elif suffix == ".csv":
            self.text = self._read_csv()
        elif suffix == ".docx":
            self.text = self._read_docx()
        elif suffix == ".rtf":
            self.text = self._read_rtf()
        else:
            raise ValueError(f"Формат {suffix} не поддерживается.")
        return self.text

    def _read_pdf(self):
        if not FITZ_AVAILABLE:
            return "Обработка PDF недоступна (требуется PyMuPDF)"

        text = ""
        with fitz.open(self.filepath) as doc:
            for page in doc:
                text += page.get_text()
        return text

    def _read_json(self):
        with open(self.filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
        return json.dumps(data, ensure_ascii=False, indent=2)

    def _read_csv(self):
        df = pd.read_csv(self.filepath)
        return df.to_string(index=False)

    def _read_docx(self):
        doc = docx.Document(self.filepath)
        text_parts = []

        # читаем параграфы
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # читаем таблицы
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    def _read_rtf(self):
        with open(self.filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return rtf_to_text(content)


def extract_job_title(text: str) -> str:
    """Извлечение названия вакансии"""
    patterns = [
        r'(?i)должность[:\s]*([^\n]+)',
        r'(?i)должность и зарплата[:\s]*([^\n]+)',
        r'(?i)позиция[:\s]*([^\n]+)',
        r'(?i)вакансия[:\s]*([^\n]+)',
        r'(?i)ищем\s+([^\n]+)',
        r'(?i)требуется\s+([^\n]+)',
        r'(?i)название[:\s]*([^\n]+)',
        r'(?i)job\s+title[:\s]*([^\n]+)'
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            title = match.group(1).strip()
            # Очищаем от лишних символов
            title = re.sub(r'[^\w\sа-яА-ЯёЁ/-]', '', title)
            return title

    # Если не нашли в шаблонах, берем первую строку
    first_line = text.split('\n')[0].strip()
    if len(first_line) < 100:  # Не слишком длинная строка
        return first_line

    return "Название вакансии не указано"


def get_embedding(text, model_path):
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    model = AutoModel.from_pretrained(model_path)

    # Добавьте эту строку - указываем устройство
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = model.to(device)

    inputs = tokenizer(text, return_tensors="pt", truncation=True, padding=True, max_length=512)

    # Перенесите inputs на то же устройство
    inputs = {k: v.to(device) for k, v in inputs.items()}

    with torch.no_grad():
        outputs = model(**inputs)
        embeddings = outputs.last_hidden_state.mean(dim=1)
        embeddings = F.normalize(embeddings, p=2, dim=1)

    # Верните на CPU для совместимости
    return embeddings.cpu()


def _generate_recommendation(score: float):
    if score >= 85.5:

        return "Хорошее соответствие. Рекомендуем пригласить на собеседование."

    else:
        return "Низкое соответствие."


if __name__ == "__main__":
    # Чтение резюме и вакансии из файлов
    resume_text = DocumentReader("data/resume/resume.rtf").extract_text()
    job_text = DocumentReader("data/job_decription/job_description.docx").extract_text()
    job_title = extract_job_title(job_text)
    print(job_title)
    model_path = "model"

    resume_emb = get_embedding(resume_text, model_path)
    job_emb = get_embedding(job_text, model_path)
    similarity = torch.mm(resume_emb, job_emb.T).item() * 100
    print(f"Cхожесть: {similarity:.2f}%")
    print(_generate_recommendation(similarity))
